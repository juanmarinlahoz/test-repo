from openpyxl import load_workbook
import os, sys
import docx as dx #para usar los docx

# ESTE PROGRAMA TIENE QUE TENER TODAS LAS FUNCIONES PARA SACAR DATOS DE LOS WORD, DE LOS XLSX, PASARLOS...
#   UNO PARA TODO
#Este programa deberia ser el que utilice para sacar datos de los .docx y meterlos en los .xlsx
#Hace falta convertir el codigo a funciones, de tal manera que sea mucho mas versatil
#

"""
Este programa pretente tener todas las funciones necesarias para extraer los datos de los archivos .xlsx generados en psychopy. 
Por el momento solo saca los datos de la "ultima" sheet del archivo, la del trial que no da feedback.
TO DO:
    -general:
        -sistema que permita separar cualquier xlsx erroneo de los otros
        -lo mismo con los incompletos
        -sistema que permita separar los tests realizados con hiragana de los otros, por ejemplo thai
        -la forma de devolver los datos
        -funcion main y manera de importar esto en otro programa
    -trial final:
        -una funcion mas general que permita extraer las respuestas a cada letra, si estan a la izquierda o a la derecha, los tiempos de respuesta...
    -no probabilistico (determinista)
        -falta todo:identificar los sheets, errores y aciertos con cada pareja, tiempos de respuesta, izquierda y derecha
    -aprendizaje no probabilistico:
        -falta todo: identificar los sheets, errotes y aciertos en cada pareja, tiempos de respuesta en cada pareja, izquierda y derecha, tiempo que tarda en aprender...
    -VERSIONES DE ARCHIVOS>
        -hay una version inicial en la que el aprendizaje probabilistico era aleatorio en bloques de 60 hasta 5.
        -la version posterior tiene un unico bloque de aprendizaje probabilisitco pseudoaleatorio (esta predeterminado siempre igual) de 300
    

"""







#   -----------------
#   ====GLOBALS======
#   -----------------


PATH=os.getcwd()#path actual
LISTAARCHIVOS=os.listdir(PATH+"/data")#hace una lista con todos los archivos en el subdirectorio data
WB=load_workbook("C:/pyzo2014a/My_programs/ARC/data/base1.xlsx") #carga la base en excel
WS=WB.active#seleciona el sheet
LISTAWbARC=os.listdir(PATH+"/ARCFdata")






"""
El megadiccionario. Contiene los nombres de todas las variables(o al menos los contendra) que se pueden sacar del docx. Contiene las instrucciones para encontrarlas: para cada variable tiene la string que la precede, la fila en la que esta y la celda en la que esta.
"""
VARDIC={
        'NOMBRE':('Nombre:',2,0),
        'NHC':("NHC:",2,0),
        'DAA':("Dosis agonistas:",2,0),
        'DIAGN':("PD de novo:",2,0),
        'DATE':("Fecha:",0,0),
        'DEBUT':("Fecha:",2,0),
        'EVOLUTION':("Años evolución:",2,0),
        'DATE':("Fecha:",0,0),
        'H&Y':("H&Y:",2,0),
        'S&E':("S&E:",2,0),
        'RBD-sq':,
        'QUIP-GAM':,
        'QUIP-SEX':,
        'QUIP-BUY':,
        'QUIP-EAT':,
        'QUIP-OTH':,        """VALORAR PONER OTRAS VARIABLES CON CADA UNO DE LOS TIPOS"""
        'QUIP-MED':,
        'LATERALIDAD':,
        'PREDOMINIO':,
        'DOSIS_LDOPA':,
        'DOSIS_PRAMIPEXOL':,
        'DOSIS_ROTIGOTINA':,
        'DOSIS_ROPIRINOL':,
        'DOSIS_RASAGILINA':,
        'DOSIS_AMANTADINA':
        
        }


#   ----------------------
#   === FUNCIONES XLSX ===
#   ----------------------

    
def search_xlsx_file(nhc,path=PATH+"/ARCFdata",extension=".xlsx"): 
    """esta funcion busca el archivo para un numero de historia dado. Para ello hace una lista con todos los nombres de archivo y luego busca el archivo. Devuelve una lista con todos los archivos que coinciden"""
    nhc=str(nhc)
    listaarchivos=os.listdir(path)
    archivos_encontrados=[]
    quitar=()
    for archivo in listaarchivos:
        if not archivo.endswith(extension):
            quitar.append(archivo)
    for archivo in quitar:
        try:
            listaarchivos.remove(archivo)
        except:
            pass
    for archivo in listaarchivos:
        if nhc in archivo:
            archivos_encontrados.append(os.path.abspath(os.path.join(path,archivo)))
    return archivos_encontrados        
            
    




"""
    ESTAS FUNCIONES SON GENERICAS PARA CUALQUIER XLSX QUE SE USE COMO BASE DE PACIENTES
"""


def selector_columna(variable,ws=WS):
    """
    Devuelve la columna cuya primera linea tiene el nombre de una variable
    """
    for number,cell in enumerate(ws.rows[0]):
        if str(cell.value)==variable:
            return number
    else:
        for number,cell in enumerate(ws.rows[0]):
            if str(cell.value)==str(variable).upper():
                return number

    
def selector_fila(valor, columna, ws=WS):
    """
    Devuelve la fila de una columna en la que hay un valor. Solo devuelve la primera, es util para NHC u otro identificador univoco
    """
    for number,cell in enumerate(ws.columns[columna]):
        if str(valor)==str(cell.value):
            return number
        

def selector_celda(nhc,variable,ws=WS):
    """
    Devuelve la celda en la en la que se encuentra el valor de una variable para un determinado paciente(NHC)
    """
    
    nhc_col=selector_columna("NHC",ws=WS)
    patient_row=selector_fila(nhc,nhc_col,ws)
    variable_col=selector_columna(variable,ws)
    return ws.cell(row=patient_row+1,column=variable_col+1)
    
    
def escribe_valor(valor,nhc,variable,ws=WS):
    """
    Escribe el valor de una variable para un paciente (NHC)
    """
    selector_celda(nhc,variable,ws).value=valor

def lee_valor(nhc,variable,ws=WS):
    """
    Da el valor de una determinada variable en un paciente
    """
    return selector_celda(nhc,variable,ws).value


def celda_vacia(nhs,variable,ws=WS):
    """
    Comprueba si el valor de una variable para un pacientes esta vacio
    """
    return lee_valor(nhc,variable,ws)==None

"""
    ESTAS FUNCIONES SON ESPECIFICAS DE LOS XLSX GENERADOS POR PSYCHOPY PARA EL ARC
"""

def select_trial_sheet(wb): 
    """esta funcion selecciona la ulitma sheet que es la que tiene las letras nuevas (G y H)"""
    for sheet in wb.worksheets:
        if sheet['E8'].value=="AH":
            return sheet
"""            
def bool_hiragana (ws):
    w   
"""             
            
def por_recompensa(ws):
    """cuenta el numero de respuestas en las que se ha elegido el estimulo A correctamente excluyendo las que habia que evitar B"""
    counter=0
    for a in range (1,len(ws.rows)):
        n=str(a)
        check=str(ws["E"+n].value).startswith("A") and not str(ws["E"+n].value).endswith("B") 
        check2=str(ws["E"+n].value).endswith("A") and not str(ws["E"+n].value).startswith("B") 
        if check:
            if "left" in ws["W"+n].value:
                counter+=1
            if  "left" in ws["X"+n].value:
                counter+=1    
        if check2:
            if "right" in ws["W"+n].value:
                counter+=1
            if  "right" in ws["X"+n].value:
                counter+=1    
    return counter

def por_castigo(ws):
    """cuenta el numero de respuestas en las que se ha evitado el estimulo B correctamente excluyendo las que habia que elegir A"""    
    counter=0
    for a in range (1,len(ws.rows)):
        n=str(a)
        check=str(ws["E"+n].value).startswith("B") and not str(ws["E"+n].value).endswith("A") 
        check2=str(ws["E"+n].value).endswith("B") and not str(ws["E"+n].value).startswith("A") 
        if check:
            if "right" in ws["W"+n].value:
                counter+=1
            if  "right" in ws["X"+n].value:
                counter+=1    
        if check2:
            if "left" in ws["W"+n].value:
                counter+=1
            if  "left" in ws["X"+n].value:
                counter+=1    
    return counter            

def rec_cas(archivo):
    """abre el archivo que le han pasado, intenta seleccionar la ultima sheet, la de los trials y devuelve el nombre de archivo y los resultados por castigo y por recompensa. El archivo se pasa como una lista. Dado que un mismo paciente puede haber realizado el ARC mas de una vez y no haberlo acabado, solo usara las versiones que tengan trials. Si hay varias versiones, solo mirara la primera que encuentre"""
    if type(archivo)==list:
        for version in archivo:
            if ".xlsx" in version:
                wb=load_workbook(version)
                try:#si no puede encontrar una sheet que tenga el trial final no da error pero no devuelve nada
                    ws=select_trial_sheet(wb)
                    return version, "\t por recompensa=" , por_recompensa(ws),"\t por castigo=", por_castigo(ws)
                except:
                    pass
    elif type(archivo)==str and ".xlsx" in archivo:
        wb=load_workbook(archivo)
        try:#si no puede encontrar una sheet que tenga el trial final no da error pero no devuelve nada
            ws=select_trial_sheet(wb)
            return archivo, "\t por recompensa=" , por_recompensa(ws),"\t por castigo=", por_castigo(ws)
        except:
           pass
    


def newones(ws):
    """suma las veces que se seleccionan estimulos nuevos, como estimacion de novelty seeking"""
    counter=0
    for a in range (1,len(ws.rows)):
        n=str(a)
        check=str(ws["E"+n].value).startswith("G") or str(ws["E"+n].value).startswith("H") 
        check2=str(ws["E"+n].value).endswith("G") or str(ws["E"+n].value).endswith("H")
        if check:
            if "left" in ws["W"+n].value:
                counter+=1
            if  "left" in ws["X"+n].value:
                counter+=1    
        if check2:
            if "right" in ws["W"+n].value:
                counter+=1
            if  "right" in ws["X"+n].value:
                counter+=1    
    return counter





#   ----------------------
#   === FUNCIONES DOCX ===
#   ----------------------    




def texto_de_docx(archivo):
    """
    Esta funcion pilla un archivo docx y devuelve el texto. Cada parrafo es una string en una lista.
    Requiere que le pasen el nombre del archivo con la path completa
    """
    try:
        f=dx.Document(archivo) #lee el archivo
    except:
        return "Este no es un archivo docx correcto" #si no es legible se queja
    texto=[] #genera una lista en la que podre meter el texto
    
    for a in f.paragraphs: #pasa por todos los parrafos del documento (los que no estan metidos en tablas)
        texto.append(a.text)#lee su texto y lo pone en la lista de texto
            
    for b in f.tables: #para cada tabla en el documento
    
        for c in b.rows:#lee las filas
        
            for d in c.cells: #lee las cells de cada fila
                for e in d.paragraphs:#lee cada parrafo de cada cell
                    
                    texto.append(e.text) #lee el texto de cada parrafo y lo mete en la lista
    return texto
    

def var_de_docx (adocx,variable):
    texto=[]
    link=VARDIC[variable]
    for p in adocx.tables[0].rows[link[1]].cells[link[2]].paragraphs:
        texto.append(p.text)
    for trozo in texto:
        #print(trozo)
        if link[0] in trozo:
            print (trozo.lstrip(link[0]))
   
   
   
    
    
"""
def texto_a_ws(crd,ws=WS):
    texto=CRD(crd).texto
    for trozo in  
"""
class CRD:
    """
    Esta clase corresponte a un archivo CRD, son archivos docx con los datos de cada participante.
    Cuando se genera un objeto de esta clase extrae el texto del archivo. La unica funcion implementada por el momento es busar una string que
    devuelve la linea entera en la que se encuentra
    """
    def __init__(self,archivo):
        self.archivo= archivo
        self.texto=texto_de_docx(archivo)
        for trozo in self.texto: #lee cada linea y saca datos
            if "Nombre" in trozo:
                self.nombre=trozo.lstrip("Nombre: ")
            if "NHC" in trozo:
                self.nhc=str(int(trozo.lstrip("NHC: ")))
            if "RBD single question:" in trozo:
                self.rbd_str=trozo.lstrip("RBD single question:")
                if "si" in self.rbd_str.lower() or "1" in self.rbd_str:
                    self.rbd= True
                else:
                    self.rbd=False


    def buscar_texto(self,cadena):
        encontrados=[]
        for linea in self.texto:
            if cadena in linea:
                encontrados.append(linea)
        if len (encontrados)>0:
            return encontrados        
        else:
            return None
        




for nombrearchivo in LISTAARCHIVOS: #selecciona solo los que sean docx. luego muestra el nombre, pero hay que sustituir la linea por otra cosa
    if ".docx" in nombrearchivo:
        texto=texto_de_docx(PATH+'/data/'+nombrearchivo)
        #print (" ".join(texto))

        for trozo in texto: #lee cada linea y saca datos
            if "Nombre" in trozo:
                nombre=trozo.lstrip("Nombre: ")
            if "NHC" in trozo:
                nhc=trozo.lstrip("NHC: ")
            if "RBD single question:  ":
                rbd_str=trozo.lstrip("RBD single question:  ")
                if "si" in rbd_str.lower():
                    rbd= True
                else:
                    rbd=False
        
        print ("Nombre: ", nombre, "\t NHC: ", nhc, "\t RBD:", rbd) #muestra los datos extraidos
        
"""        
        for cell in ws.columns[2]:#comprueba todas las casillas de la columna 2 (la de los numeros de historia) para buscar la fila en la que       esta el paciente del word
            try:
                if cell.value==int(nhc):
                    #print('el paciente', str(nombre),'con numero de historia',str(nhc),'esta en la fila:', str(cell.row))
                    fila=cell.row
                    
            except:
                pass
                
        



for col in range(1,100): #busca el numero de columna que tiene como titulo "CASADO", guarda el numero en una variable
    col_name=ws.cell(row=1,column=col).value
    if col_name== "CASADO":
        casado=col
    #print(ws.cell(row=1,column=col).value)

"""





    
"""
for archivo in LISTAWB:
    if ".xlsx" in archivo:
        print(rec_cas(PATH+"/ARCFdata/"+archivo))
            

#print(search_file(136823,PATH+"/ARCFdata"))

print(rec_cas(search_file(966147,PATH+"/ARCFdata")))
"""
